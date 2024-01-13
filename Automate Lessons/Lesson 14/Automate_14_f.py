import docx, os

os.chdir(r'C:\Users\David Huang\Downloads\Python')

d = docx.Document()
d.add_paragraph('Hello, this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('demo4.docx')
p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True
d.save('demo5.docx')

