import docx, os

os.chdir(r'C:\Users\David Huang\Downloads\Python')

d = docx.Document('demo.docx')
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]
print(p.runs)
# a run occurs every time a change in formatting occurs
print(p.runs[0].text)
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'

d.save('demo2.docx')

p.style = 'Title'
d.save('demo3.docx')
